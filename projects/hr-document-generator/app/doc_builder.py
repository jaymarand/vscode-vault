"""Generate DOCX documents by filling original Goodwill HR templates.

Templates are stored in app/templates/ and are copies of the actual
documents used at Ohio Valley Goodwill. This preserves exact formatting,
fonts, styles, and layout.
"""

from __future__ import annotations

import io
import os
import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

TEMPLATE_DIR = Path(__file__).parent / "templates"

RATING_COLS = {"E": 1, "AB": 2, "G": 3, "F": 4, "U": 5}


def _to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _replace_in_paragraph(paragraph, old: str, new: str):
    """Replace text in a paragraph while preserving formatting of the first run."""
    full = paragraph.text
    if old not in full:
        return False
    # If there's only one run or the text is all in one run, simple replace
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Text spans multiple runs — rebuild with first run's formatting
    new_text = full.replace(old, new)
    if paragraph.runs:
        first_run = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run.text = ""
        first_run.text = new_text
    return True


def _replace_in_cell(cell, old: str, new: str):
    """Replace text in a table cell."""
    for p in cell.paragraphs:
        _replace_in_paragraph(p, old, new)


def _set_paragraph_text(paragraph, text: str):
    """Set paragraph text, keeping formatting of the first run."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def _find_paragraph_containing(doc, text: str):
    """Find first paragraph containing the given text."""
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            return i, p
    return None, None


# ---------------------------------------------------------------------------
# Coaching Form
# ---------------------------------------------------------------------------

def build_coaching_docx(
    employee_name: str,
    date: str,
    location: str,
    categories: list[str],
    ai_content: dict,
    manager_name: str = "Jason Cole",
) -> bytes:
    """Fill the coaching template with provided data."""
    doc = Document(str(TEMPLATE_DIR / "coaching_template.docx"))

    # Fill employee info in table
    table = doc.tables[0]
    # Row 1 col 0: Employee Name
    _replace_in_cell(table.rows[1].cells[0], "{{PLACEHOLDER}}", employee_name)
    # Row 2 col 0: Date
    _replace_in_cell(table.rows[2].cells[0], "{{PLACEHOLDER}}", date)
    # Row 3 col 0: Location
    _replace_in_cell(table.rows[3].cells[0], "{{PLACEHOLDER}}", location)

    # Update checkboxes in purpose cell
    purpose_cell = table.rows[1].cells[1]
    for p in purpose_cell.paragraphs:
        for cat in ["Performance Improvement", "Policy Violation", "Attendance", "Other"]:
            if cat in p.text:
                check = "\u2611" if cat in categories else "\u2610"
                for run in p.runs:
                    if "\u2610" in run.text:
                        run.text = run.text.replace("\u2610", check)

    # Fill AI content sections
    replacements = {
        "{{DESCRIPTION_OF_BEHAVIOR_ISSUE}}": ai_content.get("DESCRIPTION OF BEHAVIOR/ISSUE", ""),
        "{{EXPECTED_BEHAVIOR_PERFORMANCE}}": ai_content.get("EXPECTED BEHAVIOR/PERFORMANCE", ""),
        "{{WHY_DOES_IT_MATTER}}": ai_content.get("WHY DOES IT MATTER", ""),
        "{{STEPS_TO_IMPROVE}}": ai_content.get("STEPS EMPLOYEE WILL TAKE TO IMPROVE", ""),
        "{{RESOURCES_SUPPORT}}": ai_content.get("RESOURCES OR SUPPORT PROVIDED BY MANAGER", ""),
        "{{TIMELINE}}": ai_content.get("TIMELINE FOR IMPROVEMENT", ""),
        "{{CHECKIN_DATE}}": ai_content.get("NEXT CHECK-IN DATE", ""),
        "{{WHAT_REVIEWED}}": ai_content.get("WHAT WILL BE REVIEWED", ""),
        "{{EMPLOYEE_NAME}}": employee_name,
        "{{MANAGER_NAME}}": manager_name,
    }
    for p in doc.paragraphs:
        for old, new in replacements.items():
            if old in p.text:
                _replace_in_paragraph(p, old, new)

    return _to_bytes(doc)


# ---------------------------------------------------------------------------
# Warning Record
# ---------------------------------------------------------------------------

def build_warning_docx(
    employee_name: str,
    job_title: str,
    department: str,
    date: str,
    ai_content: dict,
) -> bytes:
    """Fill the warning template with provided data."""
    doc = Document(str(TEMPLATE_DIR / "warning_template.docx"))

    # Fill employee info table (Table 0)
    table = doc.tables[0]
    # Row 0: Employee Name | ... | Vanessa Sevier | ... | Date | value
    # Find cells containing template data and replace
    for row in table.rows:
        for cell in row.cells:
            if "Vanessa Sevier" in cell.text:
                _replace_in_cell(cell, "Vanessa Sevier", employee_name)
            if "2/10/2026" in cell.text:
                _replace_in_cell(cell, "2/10/2026", date)
            if "ShopGoodwill" in cell.text:
                _replace_in_cell(cell, "ShopGoodwill", department)
            if "Ecommerce Manager" in cell.text:
                _replace_in_cell(cell, "Ecommerce Manager", job_title)

    # Replace paragraph content
    # P[4]: Nature of Violation line
    nature = ai_content.get("NATURE OF VIOLATION", "")
    additional = ai_content.get("ADDITIONAL COMMENTS", "")
    improvements = ai_content.get("REQUIRED IMPROVEMENTS", "")
    review_period = ai_content.get("REVIEW PERIOD", "30-Day Review Period")

    # Find and replace the Nature of Violation paragraph
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Nature of Violation:"):
            _set_paragraph_text(p, f"Nature of Violation: {nature}")
        elif p.text.startswith("Additional Comments:"):
            _set_paragraph_text(p, f"Additional Comments: {additional}")
        elif "Required Improvement Expectations" in p.text:
            _set_paragraph_text(p, f"Required Improvement Expectations ({review_period}):\nEffective immediately, you are expected to:")

    # Replace the body paragraphs (P[5] through P[12] are the old narrative)
    # And P[15]-P[21] are the old improvement bullets
    # Strategy: clear old content paragraphs, put new content in the first one

    # Find paragraph indices for the narrative section
    narrative_start = None
    narrative_end = None
    bullets_start = None
    bullets_end = None

    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Additional Comments:"):
            narrative_start = i
        elif "Required Improvement Expectations" in p.text:
            narrative_end = i
            bullets_start = i + 1
        elif p.text == "Performance will be monitored throughout the review period.":
            bullets_end = i

    # Clear old narrative paragraphs (after "Additional Comments" up to "Required Improvement")
    if narrative_start is not None and narrative_end is not None:
        # The Additional Comments paragraph already has the new text
        # Clear paragraphs between it and Required Improvements
        for i in range(narrative_start + 1, narrative_end):
            _set_paragraph_text(doc.paragraphs[i], "")

    # Clear old bullet paragraphs and insert new ones
    if bullets_start is not None and bullets_end is not None:
        improvement_lines = [
            line.strip().lstrip("•-0123456789. ")
            for line in improvements.strip().split("\n")
            if line.strip()
        ]
        for i in range(bullets_start, bullets_end):
            idx = i - bullets_start
            if idx < len(improvement_lines):
                _set_paragraph_text(doc.paragraphs[i], improvement_lines[idx])
            else:
                _set_paragraph_text(doc.paragraphs[i], "")

    # Replace dates in signature area
    for p in doc.paragraphs:
        if "2/10/26" in p.text:
            _replace_in_paragraph(p, "2/10/26", date)

    return _to_bytes(doc)


# ---------------------------------------------------------------------------
# Annual Review
# ---------------------------------------------------------------------------

def build_annual_review_docx(
    ai_content: dict,
    supervisor_name: str = "Jason Cole",
    template_bytes: bytes | None = None,
    employee_name: str = "",
    job_title: str = "",
    department: str = "",
    date_of_hire: str = "",
    period_from: str = "",
    period_to: str = "",
    current_pay: str = "",
    percent_increase: str = "",
    new_pay: str = "",
) -> bytes:
    """Fill an annual review template with AI-generated content.

    If template_bytes is provided (uploaded file), uses that as the base —
    employee info is already in the file. Otherwise falls back to the
    static template and replaces placeholder values.
    """
    if template_bytes:
        doc = Document(io.BytesIO(template_bytes))
    else:
        doc = Document(str(TEMPLATE_DIR / "review_template.docx"))
        # Replace static template placeholder values
        _replace_static_review_template(
            doc, employee_name, job_title, department,
            date_of_hire, period_from, period_to,
            current_pay, percent_increase, new_pay,
        )

    # Update wage info if percent_increase provided (even with uploaded template)
    if percent_increase and new_pay and len(doc.tables) > 1:
        t1 = doc.tables[1]
        for row in t1.rows:
            for cell in row.cells:
                # Fill in new pay rate if cell has empty "$"
                if cell.text.strip() == "$" and new_pay:
                    _set_paragraph_text(cell.paragraphs[0], new_pay)
                # Fill in percent if cell is empty and in the right row
            row_text = " ".join(c.text for c in row.cells)
            if "Percent of Increase:" in row_text:
                for cell in row.cells:
                    if cell.text.strip() == "" or cell.text.strip() == "$":
                        pass  # skip
                    if cell.text.strip() == "$":
                        _set_paragraph_text(cell.paragraphs[0], new_pay)
                if not any("%" in c.text for c in row.cells):
                    # Find empty cells after label and fill
                    found_label = False
                    for cell in row.cells:
                        if "Percent of Increase" in cell.text:
                            found_label = True
                            continue
                        if found_label and not cell.text.strip():
                            _set_paragraph_text(cell.paragraphs[0], f"{percent_increase}%")
                            break

    # --- Table 7 (or last table with 7 cols): Trait ratings ---
    trait_table = None
    for t in doc.tables:
        if len(t.columns) == 7 and len(t.rows) > 10:
            # Check if header matches trait rating table
            header_text = " ".join(c.text.strip() for c in t.rows[0].cells)
            if "TRAIT" in header_text or "E" in header_text:
                trait_table = t
                break

    if trait_table:
        trait_ratings = ai_content.get("trait_ratings_parsed", [])
        trait_map = {}
        for tr in trait_ratings:
            key = tr["trait"].lower().strip()
            trait_map[key] = tr

        for row_idx in range(1, len(trait_table.rows)):
            row = trait_table.rows[row_idx]
            trait_name = row.cells[0].text.strip()
            if not trait_name or trait_name == "TRAIT":
                continue

            # Clear old ratings
            for col in range(1, 6):
                _set_paragraph_text(row.cells[col].paragraphs[0], "")

            # Find matching AI rating
            matched = None
            trait_lower = trait_name.lower()
            for key, val in trait_map.items():
                if (key in trait_lower or trait_lower.split("/")[0] in key
                        or key.split("/")[0] in trait_lower):
                    matched = val
                    break

            if matched:
                rating = matched["rating"].upper().strip()
                if rating in RATING_COLS:
                    col = RATING_COLS[rating]
                    _set_paragraph_text(row.cells[col].paragraphs[0], "X")
                    row.cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                comment_text = matched.get("comment", "")
                _set_paragraph_text(row.cells[6].paragraphs[0], comment_text)
                # Set comment font to 11pt so trait table spacing stays clean
                for run in row.cells[6].paragraphs[0].runs:
                    run.font.size = Pt(11)

    # --- Supervisor's Narrative + Goals ---
    # In the template, the narrative label is in P[35] and narrative text in P[36].
    # Goals go in the empty paragraphs that follow (P[37]+).
    # Find the narrative paragraph: it's the first non-empty paragraph after
    # "Supervisor's Narrative" heading.
    narrative = ai_content.get("SUPERVISOR NARRATIVE", "")
    goals_text = ai_content.get("GOALS", "")
    goals_lines = [l.strip() for l in goals_text.strip().split("\n") if l.strip()][:3]

    narrative_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "Supervisor\u2019s Narrative" in p.text or "Supervisor's Narrative" in p.text:
            if "describe employee accomplishments" in p.text or "objectives" in p.text:
                # This is the performance evaluation narrative heading, not the
                # Transfer/Promotion section. Content goes in the next paragraph.
                narrative_idx = i + 1
                break

    if narrative_idx is not None and narrative_idx < len(doc.paragraphs):
        # Set narrative text
        _set_paragraph_text(doc.paragraphs[narrative_idx], narrative)

        # Write goals into the empty paragraphs after narrative.
        # Skip one empty paragraph for spacing, then write each goal
        # with an empty paragraph between for readability.
        goal_write_idx = narrative_idx + 1
        goals_placed = 0
        for g_idx, goal in enumerate(goals_lines):
            # Skip one empty paragraph for spacing before first goal
            if goal_write_idx < len(doc.paragraphs):
                goal_write_idx += 1  # skip spacer
            if goal_write_idx < len(doc.paragraphs):
                _set_paragraph_text(doc.paragraphs[goal_write_idx], goal)
                goals_placed += 1
                goal_write_idx += 1

        # Clear any remaining old content paragraphs up to "Employee Comments:"
        employee_comments_idx = None
        for i, p in enumerate(doc.paragraphs):
            if "Employee Comments:" in p.text:
                employee_comments_idx = i
                break

        if employee_comments_idx:
            for i in range(goal_write_idx, employee_comments_idx):
                _set_paragraph_text(doc.paragraphs[i], "")

    # --- Supervisor name in signature area ---
    for p in doc.paragraphs:
        # Look for any existing supervisor name and replace
        text = p.text.strip()
        if "Supervisor" in text and "Employee" in text and len(text) > 20:
            # This is the signature line: "Supervisor    Employee    Name"
            # Replace any name after "Employee" with the employee name
            # and keep "Supervisor" label
            pass  # Signature line labels — don't modify
        if "Cappelletty" in p.text:
            _replace_in_paragraph(p, "Cappelletty, Kimra", supervisor_name)

    return _to_bytes(doc)


# ---------------------------------------------------------------------------
# Professional Development Plan
# ---------------------------------------------------------------------------

GOAL_TYPES = [
    "Personal/Interpersonal Growth",
    "Technical/Job Skills",
    "Leadership Capabilities",
    "Career Planning",
]

RESOURCES = ["EAP", "HR", "L&D", "OC", "Wellness", "Other"]


def build_pdp_docx(
    employee_name: str,
    position: str,
    location: str,
    date: str,
    ai_content: dict,
) -> bytes:
    """Fill the PDP template with AI-generated goals and action steps."""
    doc = Document(str(TEMPLATE_DIR / "pdp_template.docx"))

    # --- Table 0: Employee info ---
    t0 = doc.tables[0]
    _replace_in_cell(t0.rows[0].cells[1], "{{NAME}}", employee_name)
    _replace_in_cell(t0.rows[0].cells[3], "{{DATE}}", date)
    _replace_in_cell(t0.rows[1].cells[1], "{{POSITION}}", position)
    _replace_in_cell(t0.rows[1].cells[3], "{{LOCATION}}", location)

    # --- Ultimate goal ---
    for p in doc.paragraphs:
        if "{{ULTIMATE_GOAL}}" in p.text:
            _replace_in_paragraph(p, "{{ULTIMATE_GOAL}}",
                                  ai_content.get("ULTIMATE GOAL", ""))

    # --- Table 1: Goals (3 columns) ---
    t1 = doc.tables[1]
    goals = ai_content.get("goals_parsed", [])

    for i, goal in enumerate(goals[:3]):
        # Goal text (row 1)
        _replace_in_cell(t1.rows[1].cells[i], "{{GOAL_" + str(i+1) + "}}", goal.get("goal", ""))

        # Check the right type checkbox (row 2)
        cell = t1.rows[2].cells[i]
        goal_type = goal.get("type", "")
        for p in cell.paragraphs:
            for gt in GOAL_TYPES:
                if gt in p.text:
                    check = "\u2611" if gt.lower().strip() in goal_type.lower().strip() or goal_type.lower().strip() in gt.lower().strip() else "\u2610"
                    for run in p.runs:
                        if "\u2610" in run.text:
                            run.text = run.text.replace("\u2610", check)

        # Why text (row 4)
        _replace_in_cell(t1.rows[4].cells[i], "{{WHY_" + str(i+1) + "}}", goal.get("why", ""))

        # Check the right timeframe (row 6)
        cell = t1.rows[6].cells[i]
        timeframe = goal.get("timeframe", "").lower().strip()
        for p in cell.paragraphs:
            for tf in ["1 month", "3 months", "6 months", "1 year"]:
                if tf in p.text:
                    check = "\u2611" if tf == timeframe else "\u2610"
                    for run in p.runs:
                        if "\u2610" in run.text:
                            run.text = run.text.replace("\u2610", check)

    # --- Table 2: Action steps (page 2) ---
    t2 = doc.tables[2]
    actions = ai_content.get("actions_parsed", [])

    for i, action in enumerate(actions[:6]):
        row = t2.rows[i + 1]  # Skip header

        # Action text
        _replace_in_cell(row.cells[0], "{{ACTION_" + str(i+1) + "}}", action.get("action", ""))

        # Check goal numbers
        goal_nums = action.get("goals", "")
        cell = row.cells[1]
        for p in cell.paragraphs:
            for gn in ["1", "2", "3"]:
                if gn in p.text:
                    check = "\u2611" if gn in goal_nums else "\u2610"
                    for run in p.runs:
                        if "\u2610" in run.text:
                            run.text = run.text.replace("\u2610", check)

        # Support partner
        _replace_in_cell(row.cells[2], "{{SUPPORT_" + str(i+1) + "}}", action.get("support", ""))

        # Resource highlighting — no checkbox change needed, the template has text labels

    return _to_bytes(doc)


FILLABLE_TEMPLATE = (
    Path(__file__).resolve().parent.parent.parent.parent
    / ".claude/skills/document-recreator/references"
    / "Professional Development Plan Template_Fillable.pdf"
)


def build_pdp_fillable_pdf(
    employee_name: str,
    position: str,
    location: str,
    date: str,
    ai_content: dict,
) -> bytes:
    """Fill the official Goodwill PDP fillable PDF template — output is editable."""
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(str(FILLABLE_TEMPLATE))
    writer = PdfWriter()
    writer.append(reader)

    goals = list(ai_content.get("goals_parsed") or [])
    while len(goals) < 3:
        goals.append({})
    actions = list(ai_content.get("actions_parsed") or [])

    # ── Text fields ──────────────────────────────────────────────────────────
    fields: dict = {
        "Name":               employee_name,
        "Position":           position,
        "Work Location":      location,
        "Date Plan Submitted": date,
        "Ultimate Goal":      ai_content.get("ULTIMATE GOAL", ""),
        "Goal 1":             goals[0].get("goal", ""),
        "Goal 2":             goals[1].get("goal", ""),
        "Goal 3":             goals[2].get("goal", ""),
        "My Why":             goals[0].get("why", ""),
        "My Why_2":           goals[1].get("why", ""),
        "My Why_3":           goals[2].get("why", ""),
    }

    # Action rows & support partners
    support_keys = [
        "Support Partnero 1 o 2 o 3",
        "Support Partnero 1 o 2 o 3_2",
        "Support Partnero 1 o 2 o 3_3",
        "Support Partnero 1 o 2 o 3_4",
        "Support Partnero 1 o 2 o 3_5",
        "Support Partnero 1 o 2 o 3_6",  # only 5 in template but safe
    ]
    for i, action in enumerate(actions[:6]):
        fields[f"ActionRow{i + 1}"] = action.get("action", "")
        if i < len(support_keys):
            fields[support_keys[i]] = action.get("support", "")

    # ── Checkboxes ───────────────────────────────────────────────────────────
    def chk(condition: bool) -> str:
        return "/Yes" if condition else "/Off"

    # Goal type checkboxes
    TYPE_FIELDS = {
        "personal":     ["1Personal", "2Personal", "3Personal"],
        "technical":    ["1Tech",     "2Tech",     "3Tech"],
        "leadership":   ["1Lead",     "2Lead",     "3Lead"],
        "career":       ["1Career",   "2Career",   "3Career"],
    }
    for key, fnames in TYPE_FIELDS.items():
        for i, fname in enumerate(fnames):
            t = goals[i].get("type", "").lower()
            fields[fname] = chk(key in t)

    # Timeframe checkboxes
    TF_FIELDS = [
        ("1 month",  ["1month",   "2month",   "3month"]),
        ("3 months", ["13month",  "23month",  "33month"]),
        ("6 months", ["16 month", "26 month", "36 month"]),
        ("1 year",   ["1Year",    "2Year",    "3Year"]),
    ]
    for tf, fnames in TF_FIELDS:
        for i, fname in enumerate(fnames):
            fields[fname] = chk(goals[i].get("timeframe", "").lower() == tf)

    # Action → goal number checkboxes
    for i, action in enumerate(actions[:6]):
        gn = str(action.get("goals", ""))
        for g in [1, 2, 3]:
            fields[f"A{i + 1}G{g}"] = chk(str(g) in gn)

    # Fill all pages
    for page in writer.pages:
        writer.update_page_form_field_values(page, fields)

    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    return buf.read()


def build_pdp_pdf(
    employee_name: str,
    position: str,
    location: str,
    date: str,
    ai_content: dict,
) -> bytes:
    """Build a landscape PDF matching the Goodwill PDP format."""
    buf = io.BytesIO()

    PAGE = landscape(letter)  # 792 x 612 pts
    MARGIN = 0.45 * inch
    W = PAGE[0] - 2 * MARGIN   # usable width ~703pt
    COL = W / 3

    BLUE = colors.HexColor("#003087")
    LIGHT = colors.HexColor("#dce6f1")

    base = getSampleStyleSheet()

    def sty(name, **kw):
        return ParagraphStyle(name, parent=base["Normal"], **kw)

    title_sty  = sty("T", fontSize=14, fontName="Helvetica-Bold", textColor=BLUE, alignment=TA_CENTER)
    label_sty  = sty("L", fontSize=9,  fontName="Helvetica-Bold")
    cell_sty   = sty("C", fontSize=8,  fontName="Helvetica",  leading=11)
    hdr_sty    = sty("H", fontSize=9,  fontName="Helvetica-Bold", textColor=colors.white, alignment=TA_CENTER)
    shdr_sty   = sty("SH", fontSize=8, fontName="Helvetica-Bold", textColor=colors.white, alignment=TA_CENTER)
    foot_sty   = sty("F", fontSize=6,  fontName="Helvetica-Oblique")

    def p(text, s=None):
        s = s or cell_sty
        return Paragraph(str(text).replace("\n", "<br/>"), s)

    def checked(options: list[str], selected: str) -> str:
        sel = selected.lower().strip()
        lines = []
        for opt in options:
            mark = "[X]" if opt.lower().strip() in sel or sel in opt.lower().strip() else "[ ]"
            lines.append(f"{mark}  {opt}")
        return "<br/>".join(lines)

    goals   = (ai_content.get("goals_parsed") or [])[:3]
    while len(goals) < 3:
        goals.append({})
    actions = (ai_content.get("actions_parsed") or [])

    doc = SimpleDocTemplate(
        buf, pagesize=PAGE,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=MARGIN, bottomMargin=MARGIN,
    )
    els = []

    # ── PAGE 1 ──────────────────────────────────────────────────────────────

    els.append(p("Professional Development Plan", title_sty))
    els.append(Spacer(1, 4))

    info = Table(
        [[p(f"<b>Name:</b>  {employee_name}", cell_sty),
          p(f"<b>Date:</b>  {date}", cell_sty)],
         [p(f"<b>Position:</b>  {position}", cell_sty),
          p(f"<b>Work Location:</b>  {location}", cell_sty)]],
        colWidths=[W / 2, W / 2],
    )
    info.setStyle(TableStyle([
        ("BOX",         (0, 0), (-1, -1), 0.5, colors.black),
        ("INNERGRID",   (0, 0), (-1, -1), 0.5, colors.black),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING",(0, 0), (-1, -1), 6),
        ("TOPPADDING",  (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING",(0,0), (-1, -1), 4),
    ]))
    els.append(info)
    els.append(Spacer(1, 4))
    els.append(p(f"<b>My ultimate goal:</b>  {ai_content.get('ULTIMATE GOAL', '')}", cell_sty))
    els.append(Spacer(1, 6))

    TYPES = ["Personal/Interpersonal Growth", "Technical/Job Skills",
             "Leadership Capabilities", "Career Planning"]
    TFS   = ["1 month", "3 months", "6 months", "1 year"]

    goal_data = [
        # Header
        [p("Goal 1", hdr_sty), p("Goal 2", hdr_sty), p("Goal 3", hdr_sty)],
        # Goal text
        [p(goals[0].get("goal", "")), p(goals[1].get("goal", "")), p(goals[2].get("goal", ""))],
        # Type header
        [p("Type of Goal", shdr_sty), p("Type of Goal", shdr_sty), p("Type of Goal", shdr_sty)],
        # Type checkboxes
        [p(checked(TYPES, goals[0].get("type", ""))),
         p(checked(TYPES, goals[1].get("type", ""))),
         p(checked(TYPES, goals[2].get("type", "")))],
        # Why header
        [p("My Why", shdr_sty), p("My Why", shdr_sty), p("My Why", shdr_sty)],
        # Why text
        [p(goals[0].get("why", "")), p(goals[1].get("why", "")), p(goals[2].get("why", ""))],
        # Timeframe header
        [p("Timeframe", shdr_sty), p("Timeframe", shdr_sty), p("Timeframe", shdr_sty)],
        # Timeframe checkboxes
        [p(checked(TFS, goals[0].get("timeframe", ""))),
         p(checked(TFS, goals[1].get("timeframe", ""))),
         p(checked(TFS, goals[2].get("timeframe", "")))],
    ]

    BLUE_ROWS = [0, 2, 4, 6]
    goal_ts = [
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    for r in BLUE_ROWS:
        goal_ts.append(("BACKGROUND", (0, r), (-1, r), BLUE))

    gt = Table(goal_data, colWidths=[COL, COL, COL])
    gt.setStyle(TableStyle(goal_ts))
    els.append(gt)

    # ── PAGE 2 ──────────────────────────────────────────────────────────────
    els.append(PageBreak())
    els.append(p("Professional Development Plan", title_sty))
    els.append(Spacer(1, 6))

    AW = [W * r for r in [0.29, 0.07, 0.14, 0.14, 0.12, 0.12, 0.12]]
    RESOURCES = ["EAP", "HR", "L&D", "OC", "Wellness", "Other"]

    act_data = [[
        p("Action", shdr_sty),
        p("Goal(s)", shdr_sty),
        p("Support Partner*", shdr_sty),
        p("Resources**", shdr_sty),
        p("Begin by:", shdr_sty),
        p("Complete by:", shdr_sty),
        p("Achieved", shdr_sty),
    ]]
    for act in actions[:6]:
        gn    = str(act.get("goals", ""))
        gchecks = "<br/>".join(
            f"{'[X]' if str(n) in gn else '[ ]'}  {n}" for n in [1, 2, 3]
        )
        res   = act.get("resource", "Other")
        rchecks = "<br/>".join(
            f"{'[X]' if r == res else '[ ]'}  {r}" for r in RESOURCES
        )
        act_data.append([
            p(act.get("action", "")),
            p(gchecks),
            p(act.get("support", "")),
            p(rchecks),
            p(act.get("begin_by", "")),
            p(act.get("complete_by", "")),
            p("[ ] Yes<br/>[ ] No"),
        ])

    at = Table(act_data, colWidths=AW)
    at.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, 0), BLUE),
        ("GRID",         (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING",   (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
    ]))
    els.append(at)
    els.append(Spacer(1, 6))
    els.append(p(
        "*Support Partner = A person who provides guidance, accountability, or expertise "
        "(e.g., manager, peer, mentor, expert etc.)  "
        "**EAP = Employee Assistance Program; L&D = OVGI Learning and Development; "
        "HR = OVGI Human Resources; OC = OVGI Opportunity Center; Wellness = OVGI Health and Wellness",
        foot_sty,
    ))

    doc.build(els)
    buf.seek(0)
    return buf.read()


def _replace_static_review_template(
    doc, employee_name, job_title, department,
    date_of_hire, period_from, period_to,
    current_pay, percent_increase, new_pay,
):
    """Replace hardcoded Marlow/Dalton values in the static template."""
    # Table 0: Employee info
    t0 = doc.tables[0]
    for row in t0.rows:
        for cell in row.cells:
            if "Marlow, Dalton" in cell.text:
                _replace_in_cell(cell, "Marlow, Dalton", employee_name)
            if "019024-Fairfield Store" in cell.text:
                _replace_in_cell(cell, "019024-Fairfield Store", department)
            if "Retail Sales Associate" in cell.text:
                _replace_in_cell(cell, "Retail Sales Associate", job_title)
            if "07/10/2024" in cell.text:
                _replace_in_cell(cell, "07/10/2024", date_of_hire)

    # Table 1: Wage info
    t1 = doc.tables[1]
    for row in t1.rows:
        for cell in row.cells:
            if "09/04/2024" in cell.text:
                _replace_in_cell(cell, "09/04/2024", period_from)
            if "09/04/2025" in cell.text:
                _replace_in_cell(cell, "09/04/2025", period_to)
            if cell.text.strip() == "12":
                _set_paragraph_text(cell.paragraphs[0], current_pay)
            if "$12.36" in cell.text:
                _replace_in_cell(cell, "$12.36", new_pay)
            if "3.0%" in cell.text:
                _replace_in_cell(cell, "3.0%", f"{percent_increase}%")

    # Table 2: Next pay review
    if len(doc.tables) > 2:
        for row in doc.tables[2].rows:
            for cell in row.cells:
                if "09/04/2026" in cell.text:
                    _replace_in_cell(cell, "09/04/2026", "")

    # Table 3: Transfer info
    if len(doc.tables) > 3:
        for row in doc.tables[3].rows:
            for cell in row.cells:
                if "Retail Team Leader" in cell.text:
                    _replace_in_cell(cell, "Retail Team Leader", "")
                if "Retail Sales Associate" in cell.text:
                    _replace_in_cell(cell, "Retail Sales Associate", job_title)
                if "019024-Fairfield Store" in cell.text:
                    _replace_in_cell(cell, "019024-Fairfield Store", department)

    # Table 6: Performance rating employee info
    if len(doc.tables) > 6:
        for row in doc.tables[6].rows:
            for cell in row.cells:
                if "Marlow, Dalton" in cell.text:
                    _replace_in_cell(cell, "Marlow, Dalton", employee_name)
                if "Retail Sales Associate" in cell.text:
                    _replace_in_cell(cell, "Retail Sales Associate", job_title)
                if "09/04/2024" in cell.text:
                    _replace_in_cell(cell, "09/04/2024", period_from)
                if "09/04/2025" in cell.text:
                    _replace_in_cell(cell, "09/04/2025", period_to)
