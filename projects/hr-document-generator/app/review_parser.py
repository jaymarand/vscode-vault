"""Parse uploaded annual review DOCX files to extract employee info."""

import io
from docx import Document


def parse_review_docx(file_bytes: bytes) -> dict:
    """Extract employee info from an uploaded annual review DOCX.

    Reads Table 0 (employee info) and Table 1 (wage info) to pull:
    - employee_name, job_title, department, date_of_hire
    - period_from, period_to, current_pay, percent_increase, new_pay
    """
    doc = Document(io.BytesIO(file_bytes))
    info = {
        "employee_name": "",
        "job_title": "",
        "department": "",
        "date_of_hire": "",
        "period_from": "",
        "period_to": "",
        "current_pay": "",
        "percent_increase": "",
        "new_pay": "",
    }

    if len(doc.tables) < 2:
        return info

    # --- Table 0: Employee info ---
    # Layout: Row 0 = [Employee Name:, _, Name, _, Department:, Dept]
    #         Row 1 = [Job Title:, Title, _, _, Date of Hire:, Date]
    t0 = doc.tables[0]
    for row in t0.rows:
        cells = [c.text.strip() for c in row.cells]
        for i, cell_text in enumerate(cells):
            if cell_text == "Employee Name:" and i + 1 < len(cells):
                # Name is in subsequent non-label cells
                for j in range(i + 1, len(cells)):
                    val = cells[j].strip()
                    if val and val not in ("Employee Name:", "Department:", "Job Title:", "Date of Hire:"):
                        info["employee_name"] = val
                        break
            if cell_text.startswith("Department:") or cell_text == "Department:":
                for j in range(i + 1, len(cells)):
                    val = cells[j].strip()
                    if val and not val.startswith(("Employee", "Job", "Date")):
                        info["department"] = val
                        break
            if cell_text == "Job Title:" and i + 1 < len(cells):
                for j in range(i + 1, len(cells)):
                    val = cells[j].strip()
                    if val and val not in ("Job Title:", "Date of Hire:", "Department:"):
                        info["job_title"] = val
                        break
            if cell_text == "Date of Hire:" and i + 1 < len(cells):
                for j in range(i + 1, len(cells)):
                    val = cells[j].strip()
                    if val and val not in ("Date of Hire:",):
                        info["date_of_hire"] = val
                        break

    # --- Table 1: Wage info ---
    # Row 0: [For the Period From:, _, date, ...| To:, date, ...]
    # Row 1: [Date of Last Increase:, _, _, date | Current Pay Rate:, _, pay, ...]
    # Row 2: [New Pay Rate:, $xx, ... | Percent of Increase:, _, pct, ...]
    t1 = doc.tables[1]
    for row in t1.rows:
        cells = [c.text.strip() for c in row.cells]
        row_text = " ".join(cells)

        if "For the Period From:" in row_text:
            # Find first date-like value after the label
            for i, c in enumerate(cells):
                if c.startswith("For the Period From"):
                    continue
                if c == "To:":
                    continue
                if c == "":
                    continue
                if not info["period_from"] and _looks_like_date(c):
                    info["period_from"] = c
                elif info["period_from"] and _looks_like_date(c) and c != info["period_from"]:
                    info["period_to"] = c
                    break

        if "Current Pay Rate:" in row_text:
            for i, c in enumerate(cells):
                if c in ("Current Pay Rate:", "Date of Last Increase:", ""):
                    continue
                # Look for pay rate value (number)
                cleaned = c.replace("$", "").replace(",", "").strip()
                try:
                    float(cleaned)
                    if "Current Pay Rate:" in " ".join(cells[:i]):
                        info["current_pay"] = cleaned
                    elif not info.get("_last_increase_date"):
                        info["_last_increase_date"] = c
                except ValueError:
                    pass

            # Fallback: find numeric values after "Current Pay Rate:"
            if not info["current_pay"]:
                found_label = False
                for c in cells:
                    if "Current Pay Rate" in c:
                        found_label = True
                        continue
                    if found_label:
                        cleaned = c.replace("$", "").replace(",", "").strip()
                        try:
                            float(cleaned)
                            info["current_pay"] = cleaned
                            break
                        except ValueError:
                            pass

        if "Percent of Increase:" in row_text:
            for c in cells:
                if "%" in c and c.replace("%", "").strip():
                    info["percent_increase"] = c.replace("%", "").strip()
                    break

        if "New Pay Rate:" in row_text:
            for c in cells:
                if c.startswith("$") and len(c) > 1:
                    info["new_pay"] = c
                    break

    # Clean up internal keys
    info.pop("_last_increase_date", None)

    return info


def _looks_like_date(text: str) -> bool:
    """Check if text looks like a date (MM/DD/YYYY)."""
    text = text.strip()
    if "/" in text:
        parts = text.split("/")
        if len(parts) >= 2:
            try:
                int(parts[0])
                int(parts[1])
                return True
            except ValueError:
                pass
    return False
